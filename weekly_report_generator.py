import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.lib.units import inch
from reportlab.platypus import Table, TableStyle
from datetime import datetime
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import logging

# 只用内置中文字体，兼容所有平台
try:
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
except Exception as e:
    logging.error(f"注册字体失败: {str(e)}")
    # 尝试使用其他中文字体
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('SimSun'))
    except Exception as e:
        logging.error(f"注册备用字体失败: {str(e)}")
        # 如果都失败了，使用默认字体
        logging.warning("使用默认字体")

class WeeklyReportGenerator:
    def __init__(self, excel_path, output_path, issue, date_str):
        self.excel_path = excel_path
        self.output_path = output_path
        self.issue = issue
        self.date_str = date_str
        self.data = None
        self.styles = getSampleStyleSheet()
        self._setup_styles()
        
    def _setup_styles(self):
        """设置PDF样式"""
        try:
            # 标题：红色、加粗、居中、较大字号
            self.styles.add(ParagraphStyle(
                name='ChineseTitle',
                fontName='STSong-Light',
                fontSize=20,
                leading=28,
                alignment=1,  # 居中
                textColor=colors.red,
                spaceAfter=10,
                spaceBefore=10,
                bold=True
            ))
            # 副标题：黑色、居中
            self.styles.add(ParagraphStyle(
                name='ChineseSubtitle',
                fontName='STSong-Light',
                fontSize=14,
                leading=20,
                alignment=1,  # 居中
                textColor=colors.black,
                spaceAfter=10
            ))
            # 一级标题：黑色、加粗、左对齐
            self.styles.add(ParagraphStyle(
                name='ChineseHeading1',
                fontName='STSong-Light',
                fontSize=13,
                leading=18,
                alignment=0,  # 左对齐
                textColor=colors.black,
                spaceBefore=10,
                spaceAfter=6,
                bold=True
            ))
            # 加粗样式
            self.styles.add(ParagraphStyle(
                name='ChineseBold',
                fontName='STSong-Light',
                fontSize=11,
                leading=18,
                alignment=0,
                textColor=colors.black,
                spaceAfter=3,
                spaceBefore=3,
                bold=True
            ))
            # 正文：黑色、常规、首行缩进
            self.styles.add(ParagraphStyle(
                name='ChineseContent',
                fontName='STSong-Light',
                fontSize=11,
                leading=18,
                alignment=0,
                firstLineIndent=24,
                textColor=colors.black,
                spaceAfter=3
            ))
            # 列表项：无缩进
            self.styles.add(ParagraphStyle(
                name='ChineseList',
                fontName='STSong-Light',
                fontSize=11,
                leading=18,
                alignment=0,
                leftIndent=12,
                textColor=colors.black,
                spaceAfter=2
            ))
            self.styles.add(ParagraphStyle(
                name='Header',
                fontName='STSong-Light',
                fontSize=9,
                alignment=1
            ))
            self.styles.add(ParagraphStyle(
                name='Footer',
                fontName='STSong-Light',
                fontSize=9,
                alignment=1
            ))
        except Exception as e:
            logging.error(f"设置样式失败: {str(e)}")
            # 使用默认样式
            self.styles.add(ParagraphStyle(
                name='ChineseTitle',
                fontSize=20,
                leading=28,
                alignment=1,
                textColor=colors.red,
                spaceAfter=10,
                spaceBefore=10,
                bold=True
            ))
    
    def load_excel_data(self):
        """加载Excel数据"""
        self.data = pd.read_excel(self.excel_path)
        self._preprocess_data()
    
    def _preprocess_data(self):
        """数据预处理"""
        # 将工作内容和计划拆分为列表
        def process_content(content):
            if isinstance(content, str):
                return [item.strip() for item in content.split('\n') if item.strip()]
            elif isinstance(content, list):
                return [item.strip() for item in content if item.strip()]
            return []
            
        self.data['last_week_work'] = self.data['上周三至本周二工作内容'].apply(process_content)
        self.data['next_week_plan'] = self.data['本周三至下周二工作计划'].apply(process_content)
        
        # 按工作类型分组
        self.grouped_data = {
            '入池': self.data[self.data['工作类型'] == '入池'],
            '入项': self.data[self.data['工作类型'] == '入项']
        }
        
        # 统计招聘数据
        self.recruitment_stats = {
            'resume': int(self.data['通过简历数量'].sum()),
            'interview': int(self.data['面试人员数量'].sum()),
            'pass': int(self.data['面试通过人员数量'].sum())
        }

        # 统计数据
        total_people = self.data['姓名'].nunique()
        pool_df = self.data[self.data['工作类型'] == '入池']
        pool_people = pool_df['姓名'].nunique()
        pool_departments = self.data['入池部门'].dropna().unique()
        pool_departments_str = '、'.join(pool_departments)
        pool_departments_count = len(pool_departments)
        project_names = self.data['项目名称'].dropna().unique()
        # 排除"其他"
        project_names = [name for name in project_names if '其他' not in str(name)]
        project_names_str = '、'.join(project_names)
        project_count = len(project_names)

    def _header_footer(self, canvas, doc):
        """添加页眉页脚"""
        canvas.saveState()
        
        # 页眉
        header = "北银金融科技有限责任公司产品研发部"
        canvas.setFont('Helvetica', 9)
        canvas.drawString(doc.leftMargin, doc.pagesize[1] - 40, header)
        
        # 页脚
        footer_text = f"页码：第{doc.page}页/共***页"
        canvas.drawString(doc.leftMargin, 30, footer_text)
        canvas.drawString(doc.leftMargin + 300, 30, "内部资料，严禁外传")
        
        canvas.restoreState()
    
    def _remove_leading_number(self, text):
        """去除开头编号（如1.、1)、1、等）"""
        return re.sub(r'^[\d一二三四五六七八九十]+[\.|、|\)|\s]+', '', str(text).strip())
    
    def generate_pdf(self):
        """生成PDF报告"""
        doc = SimpleDocTemplate(
            self.output_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        
        # 动态统计概要数据
        total_people = self.data['姓名'].nunique()
        pool_people = self.grouped_data['入池']['姓名'].nunique()
        pool_departments = self.data['入池部门'].dropna().unique()
        pool_departments_str = '、'.join(pool_departments)
        pool_departments_count = len(pool_departments)
        project_names = self.data['项目名称'].dropna().unique()
        # 排除"其他"
        project_names = [name for name in project_names if '其他' not in str(name)]
        project_names_str = '、'.join(project_names)
        project_count = len(project_names)

        # 分类"其他项目"
        def is_other_project(name):
            return '其他' in str(name)
        # 入项、入池分为"常规项目"和"其他项目"
        regular_initem = self.grouped_data['入项'][~self.grouped_data['入项']['项目名称'].apply(is_other_project)]
        other_initem = self.grouped_data['入项'][self.grouped_data['入项']['项目名称'].apply(is_other_project)]
        regular_pool = self.grouped_data['入池'][~self.grouped_data['入池']['项目名称'].apply(is_other_project)]
        other_pool = self.grouped_data['入池'][self.grouped_data['入池']['项目名称'].apply(is_other_project)]

        # 标题（两行，红色大号加粗）
        title1 = Paragraph("北银金融科技有限责任公司", ParagraphStyle(
            name='Title1', fontName='STSong-Light', fontSize=21, leading=36, alignment=1, textColor=colors.red, spaceAfter=6, spaceBefore=12, bold=True
        ))
        title2 = Paragraph("产品研发部综合业务组周例会会议纪要", ParagraphStyle(
            name='Title2', fontName='STSong-Light', fontSize=21, leading=36, alignment=1, textColor=colors.red, spaceAfter=18, bold=True
        ))
        story.append(title1)
        story.append(title2)

        # 期数（居中）
        issue_para = Paragraph(f"{datetime.now().year} 年第 {self.issue} 期", ParagraphStyle(
            name='Issue', fontName='STSong-Light', fontSize=18, alignment=1, spaceAfter=8
        ))
        story.append(issue_para)

        # 部门和日期（两列，居中）
        dept_date_table = Table(
            [[Paragraph("产品研发部", ParagraphStyle(name='Dept', fontName='STSong-Light', fontSize=16, alignment=1)),
              Paragraph(self.date_str, ParagraphStyle(name='Date', fontName='STSong-Light', fontSize=16, alignment=1))]],
            colWidths=[doc.width/2.0, doc.width/2.0]
        )
        dept_date_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(dept_date_table)

        # 分割线
        story.append(Spacer(1, 8))
        story.append(Table([['']], colWidths=[doc.width], style=TableStyle([
            ('LINEBELOW', (0,0), (-1,-1), 2, colors.black)
        ])))
        story.append(Spacer(1, 12))

        # 一、当周工作情况（加粗）
        story.append(Paragraph("一、当周工作情况", self.styles['ChineseHeading1']))
        # 动态概要段落
        summary = f"产品研发部综合业务组共计{total_people}人，组内有{pool_people}人入池{pool_departments_str}{pool_departments_count}个部门，支持行内日常工作。组内目前支持{project_count}个项目，包括{project_names_str}。"
        story.append(Paragraph(summary, self.styles['ChineseContent']))
        story.append(Paragraph("汇报详情如下：", self.styles['ChineseContent']))
        story.append(Spacer(1, 6))
        # 1.综合业务组（加粗）
        story.append(Paragraph("1.综合业务组", self.styles['ChineseBold']))
        # 1)项目进展（加粗）
        story.append(Paragraph("1)项目进展", self.styles['ChineseBold']))
        if not regular_initem.empty:
            for _, row in regular_initem.iterrows():
                project_title = f"<b>•{row['项目名称']}（{row['项目阶段']}）</b>"
                story.append(Paragraph(project_title, self.styles['ChineseBold']))
                if isinstance(row['last_week_work'], list):
                    filtered_tasks = [self._remove_leading_number(task) for task in row['last_week_work'] if task.strip()]
                    for idx, task in enumerate(filtered_tasks, 1):
                        story.append(Paragraph(f"{idx}、{task}", self.styles['ChineseList']))
                story.append(Spacer(1, 4))
        # 2)入池工作（加粗）
        story.append(Paragraph("2)入池工作", self.styles['ChineseBold']))
        story.append(Paragraph(f"目前组内有{total_people}人，{pool_people}人入池。", self.styles['ChineseContent']))
        # 过滤掉"其他"项目
        pool_no_other = self.grouped_data['入池'][~self.grouped_data['入池']['项目名称'].apply(is_other_project)]
        for dept, dept_group in pool_no_other.groupby('入池部门'):
            dept_people = dept_group['姓名'].nunique()
            story.append(Paragraph(f"•{dept}（{dept_people}人）", self.styles['ChineseBold']))
            for project_name, proj_group in dept_group.groupby('项目名称'):
                project_stage = proj_group.iloc[0]['项目阶段']
                story.append(Paragraph(f"{project_name}（{project_stage}）", self.styles['ChineseList']))
                all_tasks = []
                for _, row in proj_group.iterrows():
                    if isinstance(row['last_week_work'], list):
                        all_tasks.extend([self._remove_leading_number(task) for task in row['last_week_work'] if task.strip()])
                for idx, task in enumerate(all_tasks, 1):
                    story.append(Paragraph(f"{idx}、{task}", self.styles['ChineseList']))
            story.append(Spacer(1, 2))
        # 3)其他工作（加粗）
        story.append(Paragraph("3)其他工作", self.styles['ChineseBold']))
        other_projects = self.data[self.data['项目名称'].str.contains('其他', na=False)]
        if not other_projects.empty:
            for _, row in other_projects.iterrows():
                tasks = row['last_week_work'] if isinstance(row['last_week_work'], list) else []
                for task in tasks:
                    task = self._remove_leading_number(task)
                    if task.strip():
                        story.append(Paragraph(f"•{task}", self.styles['ChineseList']))
        # 招聘内容合并到3)其他工作
        recruitment_text = f"•招聘：简历通过{self.recruitment_stats['resume']}份，面试{self.recruitment_stats['interview']}人，通过{self.recruitment_stats['pass']}人"
        story.append(Paragraph(recruitment_text, self.styles['ChineseList']))
        story.append(Spacer(1, 4))
        
        story.append(Spacer(1, 4))
        
        # 二、下周工作计划（加粗）
        story.append(Paragraph("二、下周工作计划", self.styles['ChineseHeading1']))
        # 增加指定文案
        story.append(Paragraph("下一周产品研发部综合业务组将按计划有序推进各项目和部门入池工作，各项工作计划如下：", self.styles['ChineseContent']))
        story.append(Spacer(1, 6))
        # 1.综合业务组（加粗）
        story.append(Paragraph("1.综合业务组", self.styles['ChineseBold']))
        # 1)项目进展（加粗）
        story.append(Paragraph("1)项目进展", self.styles['ChineseBold']))
        if not regular_initem.empty:
            for _, row in regular_initem.iterrows():
                project_title = f"<b>•{row['项目名称']}（{row['项目阶段']}）</b>"
                story.append(Paragraph(project_title, self.styles['ChineseBold']))
                if isinstance(row['next_week_plan'], list):
                    filtered_plans = [self._remove_leading_number(task) for task in row['next_week_plan'] if task.strip()]
                    for idx, task in enumerate(filtered_plans, 1):
                        story.append(Paragraph(f"{idx}、{task}", self.styles['ChineseList']))
                story.append(Spacer(1, 4))
        # 2)入池工作（加粗）
        story.append(Paragraph("2)入池工作", self.styles['ChineseBold']))
        story.append(Paragraph(f"目前组内有{total_people}人，{pool_people}人入池。", self.styles['ChineseContent']))
        # 过滤掉"其他"项目
        pool_no_other = self.grouped_data['入池'][~self.grouped_data['入池']['项目名称'].apply(is_other_project)]
        for dept, dept_group in pool_no_other.groupby('入池部门'):
            dept_people = dept_group['姓名'].nunique()
            story.append(Paragraph(f"•{dept}（{dept_people}人）", self.styles['ChineseBold']))
            for project_name, proj_group in dept_group.groupby('项目名称'):
                project_stage = proj_group.iloc[0]['项目阶段']
                story.append(Paragraph(f"{project_name}（{project_stage}）", self.styles['ChineseList']))
                all_tasks = []
                for _, row in proj_group.iterrows():
                    if isinstance(row['next_week_plan'], list):
                        all_tasks.extend([self._remove_leading_number(task) for task in row['next_week_plan'] if task.strip()])
                for idx, task in enumerate(all_tasks, 1):
                    story.append(Paragraph(f"{idx}、{task}", self.styles['ChineseList']))
            story.append(Spacer(1, 2))
        # 3)其他工作（加粗）
        story.append(Paragraph("3)其他工作", self.styles['ChineseBold']))
        other_projects = self.data[self.data['项目名称'].str.contains('其他', na=False)]
        if not other_projects.empty:
            for _, row in other_projects.iterrows():
                plans = row['next_week_plan'] if isinstance(row['next_week_plan'], list) else []
                for plan in plans:
                    plan = self._remove_leading_number(plan)
                    if plan.strip():
                        story.append(Paragraph(f"•{plan}", self.styles['ChineseList']))
        # 招聘内容合并到3)其他工作
        story.append(Paragraph("•招聘：持续招聘工作", self.styles['ChineseList']))
        story.append(Spacer(1, 4))
        
        story.append(Spacer(1, 4))
        
        # 生成PDF（无页眉页脚）
        doc.build(story)
    
    def run(self):
        """运行生成器"""
        self.load_excel_data()
        self.generate_pdf()

def generate_word_report(excel_path, output_path, issue, date_str):
    import pandas as pd
    import re
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    df = pd.read_excel(excel_path)
    # 先处理工作内容字段
    def process_content(content):
        if isinstance(content, str):
            return [item.strip() for item in content.split('\n') if item.strip()]
        elif isinstance(content, list):
            return [item.strip() for item in content if item.strip()]
        return []
    df['last_week_work'] = df['上周三至本周二工作内容'].apply(process_content)
    df['next_week_plan'] = df['本周三至下周二工作计划'].apply(process_content)
    other_projects = df[df['项目名称'].str.contains('其他', na=False)]

    # 统计数据
    total_people = df['姓名'].nunique()
    pool_df = df[df['工作类型'] == '入池']
    pool_people = pool_df['姓名'].nunique()
    pool_departments = df['入池部门'].dropna().unique()
    pool_departments_str = '、'.join(pool_departments)
    pool_departments_count = len(pool_departments)
    project_names = df['项目名称'].dropna().unique()
    # 排除"其他"
    project_names = [name for name in project_names if '其他' not in str(name)]
    project_names_str = '、'.join(project_names)
    project_count = len(project_names)
    grouped_data = {
        '入池': pool_df,
        '入项': df[df['工作类型'] == '入项']
    }
    recruitment_stats = {
        'resume': int(df['通过简历数量'].sum()),
        'interview': int(df['面试人员数量'].sum()),
        'pass': int(df['面试通过人员数量'].sum())
    }
    def remove_leading_number(text):
        return re.sub(r'^[\d一二三四五六七八九十]+[\.|、|\)|\s]+', '', str(text).strip())
    def is_other_project(name):
        return '其他' in str(name)
    
    # 入项、入池分为"常规项目"和"其他项目"，此时所有DataFrame都带有last_week_work/next_week_plan
    regular_initem = grouped_data['入项'][~grouped_data['入项']['项目名称'].apply(is_other_project)]
    other_initem = grouped_data['入项'][grouped_data['入项']['项目名称'].apply(is_other_project)]
    regular_pool = grouped_data['入池'][~grouped_data['入池']['项目名称'].apply(is_other_project)]
    other_pool = grouped_data['入池'][grouped_data['入池']['项目名称'].apply(is_other_project)]

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题1（红色、加粗、居中、字号21）
    p = doc.add_paragraph()
    run = p.add_run('北银金融科技有限责任公司')
    run.font.size = Pt(21)
    run.bold = True
    run.font.color.rgb = RGBColor(0xE6, 0x19, 0x19)  # 红色
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    # 标题2（红色、加粗、居中、字号21）
    p = doc.add_paragraph()
    run = p.add_run('产品研发部综合业务组周例会会议纪要')
    run.font.size = Pt(21)
    run.bold = True
    run.font.color.rgb = RGBColor(0xE6, 0x19, 0x19)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    # 期数（居中，字号18）
    p = doc.add_paragraph()
    run = p.add_run(f'{pd.Timestamp.now().year} 年第 {issue} 期')
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    # 部门和日期（两列，居中，字号16）
    table = doc.add_table(rows=1, cols=2)
    table.alignment = 1  # 居中
    table.autofit = True
    table.allow_autofit = True
    table.columns[0].width = Cm(7)
    table.columns[1].width = Cm(7)
    cell1 = table.cell(0, 0)
    cell2 = table.cell(0, 1)
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run('产品研发部')
    run1.font.size = Pt(16)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = cell2.paragraphs[0]
    run2 = p2.add_run(date_str)
    run2.font.size = Pt(16)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 分割线（黑色粗线）
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('')
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '16')  # 粗线
    bottom.set(qn('w:color'), '000000')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    # 空行
    doc.add_paragraph()
    # 一级标题
    p = doc.add_paragraph()
    run = p.add_run('一、当周工作情况')
    run.bold = True
    run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    # 概要段落
    summary = f"产品研发部综合业务组共计{total_people}人，组内有{pool_people}人入池{pool_departments_str}{pool_departments_count}个部门，支持行内日常工作。组内目前支持{project_count}个项目，包括{project_names_str}。"
    para = doc.add_paragraph(summary)
    para.paragraph_format.first_line_indent = Cm(1)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    para = doc.add_paragraph('汇报详情如下：')
    para.paragraph_format.first_line_indent = Cm(1)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()
    # 二级标题
    p = doc.add_paragraph()
    run = p.add_run('1.综合业务组')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    # 三级标题
    p = doc.add_paragraph()
    run = p.add_run('1)项目进展')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    if not regular_initem.empty:
        for _, row in regular_initem.iterrows():
            p = doc.add_paragraph()
            run = p.add_run(f'•{row["项目名称"]}（{row["项目阶段"]}）')
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.5
    # 三级标题
    p = doc.add_paragraph()
    run = p.add_run('2)入池工作')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p = doc.add_paragraph()
    run = p.add_run(f"目前组内有{total_people}人，{pool_people}人入池。")
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.5
    pool_no_other = pool_df[~pool_df['项目名称'].apply(is_other_project)]
    for dept, dept_group in pool_no_other.groupby('入池部门'):
        dept_people = dept_group['姓名'].nunique()
        para = doc.add_paragraph(f'•{dept}（{dept_people}人）')
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.line_spacing = 1.5
        for project_name, proj_group in dept_group.groupby('项目名称'):
            project_stage = proj_group.iloc[0]['项目阶段']
            para = doc.add_paragraph()
            run = para.add_run(f'{project_name}（{project_stage}）')
            run.bold = True
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = 1.5
            all_tasks = []
            for _, row in proj_group.iterrows():
                if isinstance(row['last_week_work'], list):
                    all_tasks.extend([remove_leading_number(task) for task in row['last_week_work'] if task.strip()])
            for idx, task in enumerate(all_tasks, 1):
                para = doc.add_paragraph(f'{idx}、{task}')
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.5
    # 三级标题
    p = doc.add_paragraph()
    run = p.add_run('3)其他工作')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    if not other_projects.empty:
        for _, row in other_projects.iterrows():
            if isinstance(row['last_week_work'], list):
                tasks = [remove_leading_number(task) for task in row['last_week_work'] if task.strip()]
                for task in tasks:
                    para = doc.add_paragraph(f'•{task}')
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.space_after = Pt(1)
                    para.paragraph_format.line_spacing = 1.5
    para = doc.add_paragraph(f'•招聘：简历通过{recruitment_stats["resume"]}份，面试{recruitment_stats["interview"]}人，通过{recruitment_stats["pass"]}人')
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    # 一级标题
    p = doc.add_paragraph()
    run = p.add_run('二、下周工作计划')
    run.bold = True
    run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    para = doc.add_paragraph('下一周产品研发部综合业务组将按计划有序推进各项目和部门入池工作，各项工作计划如下：')
    para.paragraph_format.first_line_indent = Cm(1)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()
    # 二级标题
    p = doc.add_paragraph()
    run = p.add_run('1.综合业务组')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    # 三级标题
    p = doc.add_paragraph()
    run = p.add_run('1)项目进展')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    if not regular_initem.empty:
        for _, row in regular_initem.iterrows():
            p = doc.add_paragraph()
            run = p.add_run(f'•{row["项目名称"]}（{row["项目阶段"]}）')
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.5
            if isinstance(row['next_week_plan'], list):
                filtered_plans = [remove_leading_number(task) for task in row['next_week_plan'] if task.strip()]
                for idx, task in enumerate(filtered_plans, 1):
                    para = doc.add_paragraph(f'{idx}、{task}')
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.space_after = Pt(1)
                    para.paragraph_format.line_spacing = 1.5
    # 三级标题
    p = doc.add_paragraph()
    run = p.add_run('2)入池工作')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p = doc.add_paragraph()
    run = p.add_run(f"目前组内有{total_people}人，{pool_people}人入池。")
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.5
    pool_no_other = pool_df[~pool_df['项目名称'].apply(is_other_project)]
    for dept, dept_group in pool_no_other.groupby('入池部门'):
        dept_people = dept_group['姓名'].nunique()
        para = doc.add_paragraph(f'•{dept}（{dept_people}人）')
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.line_spacing = 1.5
        for project_name, proj_group in dept_group.groupby('项目名称'):
            project_stage = proj_group.iloc[0]['项目阶段']
            para = doc.add_paragraph()
            run = para.add_run(f'{project_name}（{project_stage}）')
            run.bold = True
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = 1.5
            all_tasks = []
            for _, row in proj_group.iterrows():
                if isinstance(row['next_week_plan'], list):
                    all_tasks.extend([remove_leading_number(task) for task in row['next_week_plan'] if task.strip()])
            for idx, task in enumerate(all_tasks, 1):
                para = doc.add_paragraph(f'{idx}、{task}')
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.5
    # 三级标题
    p = doc.add_paragraph()
    run = p.add_run('3)其他工作')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    other_projects = df[df['项目名称'].str.contains('其他', na=False)]
    if not other_projects.empty:
        for _, row in other_projects.iterrows():
            if isinstance(row['next_week_plan'], list):
                plans = [remove_leading_number(task) for task in row['next_week_plan'] if task.strip()]
                for plan in plans:
                    para = doc.add_paragraph(f'•{plan}')
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.space_after = Pt(1)
                    para.paragraph_format.line_spacing = 1.5
    para = doc.add_paragraph('•招聘：持续招聘工作')
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    doc.save(output_path)

if __name__ == "__main__":
    # 示例使用
    generator = WeeklyReportGenerator(
        excel_path="sample_data.xlsx",
        output_path="weekly_report.pdf",
        issue="1",
        date_str="2024年1月1日"
    )
    generator.run() 